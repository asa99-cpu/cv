import streamlit as st
import json
import docx
from docx.shared import Pt
from io import BytesIO
from weasyprint import HTML

# Load CV data from JSON
with open("data/cv_data.json", "r") as file:
    cv_data = json.load(file)

# Function to generate Word CV
def create_word_cv(data):
    doc = docx.Document()
    doc.add_heading(data["name"], level=1)

    doc.add_paragraph(f"📞 Phone: {data['phone']}")
    doc.add_paragraph(f"📧 Email: {data['email']}")
    doc.add_paragraph(f"🎂 Date of Birth: {data['dob']}")
    doc.add_paragraph(f"🏠 Nationality: {data['nationality']}")
    
    doc.add_heading("Education", level=2)
    doc.add_paragraph(data["education"])

    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(data["skills"]))

    doc.add_heading("Languages", level=2)
    doc.add_paragraph(", ".join(data["languages"]))

    doc.add_heading("Internships", level=2)
    for internship in data["internships"]:
        doc.add_paragraph(f"- {internship}")

    # Save to memory
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to generate PDF CV
def create_pdf_cv(data):
    html_content = f"""
    <h1>{data['name']}</h1>
    <p><strong>📞 Phone:</strong> {data['phone']}</p>
    <p><strong>📧 Email:</strong> {data['email']}</p>
    <p><strong>🎂 Date of Birth:</strong> {data['dob']}</p>
    <p><strong>🏠 Nationality:</strong> {data['nationality']}</p>
    <h2>Education</h2>
    <p>{data['education']}</p>
    <h2>Skills</h2>
    <p>{", ".join(data['skills'])}</p>
    <h2>Languages</h2>
    <p>{", ".join(data['languages'])}</p>
    <h2>Internships</h2>
    <ul>
        {"".join([f"<li>{internship}</li>" for internship in data['internships']])}
    </ul>
    """
    pdf_buffer = BytesIO()
    HTML(string=html_content).write_pdf(pdf_buffer)
    pdf_buffer.seek(0)
    return pdf_buffer

# Streamlit App
st.title("📄 CV Generator")

# Show CV details
st.subheader(cv_data["name"])
st.write(f"📞 **Phone:** {cv_data['phone']}")
st.write(f"📧 **Email:** {cv_data['email']}")
st.write(f"🎂 **Date of Birth:** {cv_data['dob']}")
st.write(f"🏠 **Nationality:** {cv_data['nationality']}")

st.subheader("Education")
st.write(cv_data["education"])

st.subheader("Skills")
st.write(", ".join(cv_data["skills"]))

st.subheader("Languages")
st.write(", ".join(cv_data["languages"]))

st.subheader("Internships")
for internship in cv_data["internships"]:
    st.write(f"- {internship}")

# Download buttons
st.subheader("Download CV")
word_buffer = create_word_cv(cv_data)
st.download_button("📥 Download as Word", word_buffer, file_name="cv.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

pdf_buffer = create_pdf_cv(cv_data)
st.download_button("📥 Download as PDF", pdf_buffer, file_name="cv.pdf", mime="application/pdf")
